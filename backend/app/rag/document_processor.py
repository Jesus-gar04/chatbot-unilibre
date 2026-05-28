import os
import tempfile
from datetime import datetime
from pathlib import Path
from typing import List, Tuple

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document

from app.config import settings


# ── Supabase client singleton ─────────────────────────────────────────────────

_supabase = None


def get_supabase():
    global _supabase
    if _supabase is None:
        from supabase import create_client
        if not settings.supabase_url or not settings.supabase_service_role_key:
            raise RuntimeError(
                "SUPABASE_URL y SUPABASE_SERVICE_ROLE_KEY deben estar configurados en .env"
            )
        _supabase = create_client(
            settings.supabase_url,
            settings.supabase_service_role_key,
        )
    return _supabase


# ── Text extraction ──────────────────────────────────────────────────────────

def extract_pdf(file_path: str) -> List[Document]:
    import fitz  # PyMuPDF
    docs = []
    pdf = fitz.open(file_path)
    for page_num in range(len(pdf)):
        text = pdf[page_num].get_text()
        if text.strip():
            docs.append(Document(
                page_content=text,
                metadata={"page": page_num + 1, "source": Path(file_path).name},
            ))
    pdf.close()
    return docs


def extract_docx(file_path: str) -> List[Document]:
    from docx import Document as DocxDoc
    doc = DocxDoc(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return [Document(
        page_content="\n".join(paragraphs),
        metadata={"page": 1, "source": Path(file_path).name},
    )]


def extract_txt(file_path: str) -> List[Document]:
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read()
    return [Document(
        page_content=text,
        metadata={"page": 1, "source": Path(file_path).name},
    )]


# ── Chunking ─────────────────────────────────────────────────────────────────

def process_document(file_path: str, doc_id: str) -> Tuple[List[Document], int]:
    """
    Extrae texto del archivo y lo divide en chunks con overlap.
    Agrega doc_id al metadata de cada chunk para poder borrarlo luego.
    """
    ext = Path(file_path).suffix.lower()

    if ext == ".pdf":
        raw = extract_pdf(file_path)
    elif ext in (".docx", ".doc"):
        raw = extract_docx(file_path)
    elif ext == ".txt":
        raw = extract_txt(file_path)
    else:
        raise ValueError(f"Tipo de archivo no soportado: {ext}")

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.chunk_size,
        chunk_overlap=settings.chunk_overlap,
    )
    chunks = splitter.split_documents(raw)

    for i, chunk in enumerate(chunks):
        chunk.metadata["doc_id"] = doc_id
        chunk.metadata["chunk_index"] = i

    return chunks, len(chunks)


# ── Supabase Storage ──────────────────────────────────────────────────────────

_MIME = {
    ".pdf":  "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".txt":  "text/plain",
}


def upload_to_storage(doc_id: str, ext: str, content: bytes):
    """Sube el archivo original al bucket de Supabase Storage."""
    sb = get_supabase()
    path = f"{doc_id}{ext}"
    sb.storage.from_(settings.storage_bucket).upload(
        path=path,
        file=content,
        file_options={"content-type": _MIME.get(ext, "application/octet-stream")},
    )


def delete_from_storage(doc_id: str):
    """Elimina el archivo del bucket, probando todas las extensiones posibles."""
    sb = get_supabase()
    for ext in (".pdf", ".docx", ".txt"):
        try:
            sb.storage.from_(settings.storage_bucket).remove([f"{doc_id}{ext}"])
        except Exception:
            pass


# ── Metadata en Supabase PostgreSQL ──────────────────────────────────────────

def add_document_metadata(
    doc_id: str, name: str, file_type: str, size: int, chunks: int
):
    """Inserta el registro del documento en la tabla 'documents' de Supabase."""
    sb = get_supabase()
    sb.table("documents").insert({
        "id": doc_id,
        "name": name,
        "type": file_type,
        "size": size,
        "chunks": chunks,
        "upload_date": datetime.utcnow().isoformat(),
        "status": "indexed",
    }).execute()


def remove_document_metadata(doc_id: str):
    """Elimina el registro del documento de la tabla 'documents'."""
    sb = get_supabase()
    sb.table("documents").delete().eq("id", doc_id).execute()


def get_all_documents() -> List[dict]:
    """Retorna todos los documentos indexados."""
    sb = get_supabase()
    result = sb.table("documents").select("*").order("upload_date", desc=True).execute()
    return result.data or []


def get_document(doc_id: str) -> dict | None:
    """Retorna un documento por su ID."""
    sb = get_supabase()
    result = sb.table("documents").select("*").eq("id", doc_id).execute()
    return result.data[0] if result.data else None
